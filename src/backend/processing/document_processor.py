# Extract text from PDF and .docx files, outputting JSON with metadata

import os
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

import re
import uuid # Unique identifier for documents
import PyPDF2 # PDF processor
from docx import Document # Document processor
from typing import Dict, Optional 
from src.backend.utils.logger import logging

class DocumentProcessor:
    def __init__(self):
        self.valid_extensions = {".pdf", ".docx"}

    def extract_text(self, file_path: str, doc_type: str) -> Optional[Dict]:
        """
        Extract text from PDF or .docx files and return structured JSON data.
        
        Args:
            file_path (str): Path to the document (PDF or .docx).
            doc_type (str): Type of document ('resume' or 'job').
        
        Returns:
            dict: JSON data with doc_id, type, text, word_count, file_name, and file_size_mb, or None if error.
        """
        # Validate file
        if not os.path.exists(file_path):
            logging.error(f"File not found: {file_path}")
            return None
        
        # Validate file size
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > 5:
            logging.error(f"File too large: {file_size_mb:.2f}MB (max 5MB)")
            return None
        
        # Validate file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.valid_extensions:
            logging.error(f"Unsupported file format: {file_ext}")
            return None
        
        logging.info(f"Processing {doc_type} file: {file_path}")
        try:
            # Generate unique document ID
            doc_id = str(uuid.uuid4())
            text = ""
            word_count = 0

            if file_ext == ".pdf":
                try:
                    with open(file_path, "rb") as file:
                        reader = PyPDF2.PdfReader(file)
                        if len(reader.pages) == 0:
                            logging.warning(f"Empty PDF: {file_path}")
                            return None
                        for page in reader.pages:
                            page_text = page.extract_text() or ""
                            text += page_text + " "
                except Exception as e:
                    logging.error(f"PDF parsing error for {file_path}: {str(e)}")
                    return None

            elif file_ext == ".docx":
                try:
                    doc = Document(file_path)
                    if not doc.paragraphs:
                        logging.warning(f"Empty DOCX: {file_path}")
                        return None
                    for para in doc.paragraphs:
                        text += para.text + " "
                except Exception as e:
                    logging.error(f"DOCX parsing error for {file_path}: {str(e)}")
                    return None

            # Clean text: remove extra whitespace, special characters, and non-ASCII
            text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII
            text = re.sub(r'[\n\r\t]+', ' ', text)  # Remove newlines, tabs
            text = re.sub(r'\s+', ' ', text).strip()  # Normalize whitespace
            text = re.sub(r'[•●○-]', ' ', text)  # Remove bullets
            
            if not text:
                logging.warning(f"No text extracted from {file_path}")
                return None
            
            # Calculate word count
            word_count = len(text.split())
            
            # Return structured JSON
            result = {
                "doc_id": doc_id,
                "doc_type": doc_type,
                "text": text,
                "word_count": word_count,
                "file_name": os.path.basename(file_path),
                "file_size_mb": round(file_size_mb, 2)
            }
            logging.info(f"Text extraction successful for {file_path}. Word count: {word_count}")
            return result
        
        except Exception as e:
            logging.error(f"Unexpected error processing {file_path}: {str(e)}")
            return None